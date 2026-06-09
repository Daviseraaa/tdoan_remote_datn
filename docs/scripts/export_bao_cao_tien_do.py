# -*- coding: utf-8 -*-
"""Xuất báo cáo tiến độ StationHub ra file Word."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1] / "Bao-cao-tien-do-21-04-30-05-2026.docx"


def set_doc_font(doc: Document, name: str = "Times New Roman", size: int = 13):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return h


def add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.bold = bold
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build():
    doc = Document()
    set_doc_font(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("BÁO CÁO TIẾN ĐỘ DỰ ÁN")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(
        "Nền tảng quản lý Agent, Task, Workflow & Tự động hóa (StationHub Server)\n"
        "Giai đoạn: 21/04/2026 – 30/05/2026"
    )
    r2.font.size = Pt(13)
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.add_paragraph()

    weeks = [
        {
            "title": "Tuần 1 (21/04 – 27/04/2026)",
            "goal": "Khảo sát yêu cầu, chốt phạm vi đồ án và kiến trúc tổng thể trước khi khởi tạo mã nguồn.",
            "work": [
                "Phân tích bài toán: quản lý fleet agent, giao task từ server, workflow/automation, giám sát realtime.",
                "Chốt stack: NestJS 11 + PostgreSQL (Prisma) + Redis/BullMQ + Socket.IO; React + Vite (admin); Rust + Electron (agent); Chrome Extension.",
                "Phác thảo luồng: đăng nhập admin (JWT) → agent /ws/agent → dispatch task → cập nhật trạng thái/log.",
                "Chuẩn bị môi trường: Docker (Postgres, Redis), Node 20+, Rust toolchain.",
            ],
            "result": [
                "Tài liệu ý tưởng kiến trúc monorepo (server / admin / agent).",
                "Sẵn sàng triển khai codebase (tuần sau).",
            ],
            "note": "Chưa có commit repository; giai đoạn thiết kế & chuẩn bị.",
        },
        {
            "title": "Tuần 2 (28/04 – 04/05/2026)",
            "goal": "Khởi tạo dự án end-to-end: backend, admin cơ bản, agent, schema dữ liệu.",
            "work": [
                "29/04 — init project: monorepo (~206 file). Backend (Auth, Users, Agents, Tasks, Workflow, Health, Admin); Prisma schema; docker-compose, seed; BullMQ, Pino.",
                "Admin-stationhub: SPA đăng nhập, layout, trang quản lý ban đầu.",
                "Agent: core Rust Socket.IO, desktop Electron, chrome-extension, chrome-bridge.",
                "30/04 — telegram log: logging / tích hợp Telegram (trigger, bot).",
                "30/04 — remote fix: sửa kết nối agent từ xa, ổn định WebSocket.",
            ],
            "result": [
                "Hệ thống chạy dev: API localhost:3000/api, admin Vite, agent kết nối server.",
                "Có tài khoản seed và health check.",
            ],
            "note": "Đồng bộ nhiều thành phần (Rust + Nest + extension) trong một milestone.",
        },
        {
            "title": "Tuần 3 (05/05 – 11/05/2026)",
            "goal": "Hoàn thiện nghiệp vụ lõi: task queue, workflow, CRUD admin, agent lifecycle.",
            "work": [
                "TasksService: trạng thái PENDING → RUNNING → terminal; log; hủy task.",
                "Workflow: CRUD, graph steps, execute thủ công, WorkflowRun / WorkflowStepRun.",
                "Agents: đăng ký, agentKey, heartbeat, metadata, cluster theo user.",
                "Admin: Dashboard, Agents/Tasks, protected route + refresh token.",
                "Agent Rust: tools::dispatch theo loại task.",
                "Tài liệu docs/project-memory (architecture, flows, agent).",
            ],
            "result": [
                "Luồng admin tạo task → agent nhận socket → trả kết quả ổn định nội bộ.",
                "Workflow lưu và trigger cơ bản.",
            ],
            "note": "Không có commit milestone riêng; tích lũy hướng ver1 (20/05).",
        },
        {
            "title": "Tuần 4 (12/05 – 18/05/2026)",
            "goal": "Mở rộng automation: trigger, task template, chrome/desktop recording.",
            "work": [
                "Workflow triggers: SCHEDULE, TELEGRAM; WorkflowTrigger, execution history.",
                "Task templates: Desktop Automation, Chrome Extension, Screen Capture.",
                "Chrome scripts & Desktop recordings: lưu/sync từ agent.",
                "Admin Automations (read-only), Audit log.",
                "Chuẩn hóa API types và hooks TanStack Query.",
            ],
            "result": [
                "Admin quản lý nguồn automation gắn agent.",
                "Trigger cấu hình và theo dõi lần chạy.",
            ],
            "note": "Đồng bộ script/recording agent ↔ server (insert/update/skip).",
        },
        {
            "title": "Tuần 5 (19/05 – 25/05/2026)",
            "goal": "Bản ver1; tích hợp trình duyệt; ổn định trước demo.",
            "work": [
                "20/05 — ver1 (~203 file): workflow engine, import nguồn, biến workflow; WorkflowEditor React Flow; import script/recording/template.",
                "23/05 — final1 before browser: CloakBrowser, Chrome profile, OPEN_BROWSER.",
                "24/05 — final 24-5: refactor agent desktop & extension sau tích hợp browser.",
            ],
            "result": [
                "Workflow editor đồ họa: node, edge, chạy thử.",
                "Agent task phức tạp hơn (browser profile, extension steps).",
            ],
            "note": "Debug Native Messaging + extension + Rust runner.",
        },
        {
            "title": "Tuần 6 (26/05 – 30/05/2026)",
            "goal": "Hoàn thiện UI admin (desktop + mobile), sửa UX/layout, chốt bản final.",
            "work": [
                "30/05 — final 1: responsive mobile (sidebar drawer, NavLayoutContext).",
                "Workflows: full-width, list mobile nền đặc, ẩn list khi editor, Alt chọn nhiều / chuột trái pan.",
                "Chrome Scripts / Desktop Recordings: cluster agent, editor React Flow.",
                "Task Template Wizard 3 bước, flow canvas mobile.",
                "Automations, Settings, Dashboard, Agents, Tasks: mobile.",
                "TopBar flex column (--app-topbar-height); randomId; EditorErrorBoundary.",
            ],
            "result": [
                "Admin dùng được mobile (~375px) và desktop (≥1024px).",
                "Luồng chính: login → agent → task/template → workflow → automation.",
            ],
            "note": "Milestone git: init → telegram → ver1 → browser → 24-5 → final 1.",
        },
    ]

    for w in weeks:
        add_heading(doc, w["title"], level=2)
        add_para(doc, "Mục tiêu tuần", bold=True)
        add_para(doc, w["goal"])
        add_para(doc, "Công việc đã thực hiện", bold=True)
        add_bullets(doc, w["work"])
        add_para(doc, "Kết quả", bold=True)
        add_bullets(doc, w["result"])
        if w.get("note"):
            add_para(doc, "Ghi chú", bold=True)
            add_para(doc, w["note"])
        doc.add_paragraph()

    add_heading(doc, "TỔNG HỢP GIAI ĐOẠN (21/04 – 30/05)", level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Hạng mục"
    hdr[1].text = "Trạng thái"
    rows_data = [
        ("Backend API + Auth + RBAC", "Hoàn thành lõi"),
        ("Agent WebSocket + task dispatch", "Hoàn thành lõi"),
        ("Workflow editor + execute", "Hoàn thành, tinh chỉnh UX"),
        ("Chrome / Desktop automation", "Hoàn thành tích hợp"),
        ("Telegram trigger & logging", "Đã có"),
        ("Admin SPA desktop", "Hoàn thành"),
        ("Admin SPA mobile", "Hoàn thành (30/05)"),
        ("Tài liệu kỹ thuật nội bộ", "docs/project-memory"),
    ]
    for a, b in rows_data:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b

    doc.add_paragraph()
    add_para(
        doc,
        "Milestone Git: init (29/04) → telegram log / remote fix (30/04) → ver1 (20/05) → "
        "final1 before browser (23/05) → final 24-5 (24/05) → final 1 (30/05).",
    )

    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    build()
